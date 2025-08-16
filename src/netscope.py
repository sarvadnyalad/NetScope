#!/usr/bin/env python
# NetScope: Automated Nmap Recon & Reporting
# Works on Windows, Linux, macOS (Windows-focused instructions in README)

import argparse
import ipaddress
import json
import os
import platform
import re
import subprocess
import sys
import time
from datetime import datetime
from collections import defaultdict

# Third-party deps
import nmap
from docx import Document
from docx.shared import Inches

try:
    import matplotlib
    import matplotlib.pyplot as plt
    MATPLOTLIB_OK = True
except Exception:
    MATPLOTLIB_OK = False

# -------------------------
# Utility & Environment
# -------------------------

def shell(cmd):
    """Run a command and return (code, stdout, stderr)"""
    p = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, shell=True, text=True)
    out, err = p.communicate()
    return p.returncode, out, err

def windows_ipconfig_network():
    """
    Parse ipconfig to produce a default /24 network CIDR, preferring the active adapter.
    Returns a CIDR string like '192.168.0.0/24' or None if not found.
    """
    code, out, err = shell("ipconfig")
    if code != 0:
        return None

    # Find IPv4 and Subnet Mask pairs per adapter
    adapters = []
    current = {}
    for line in out.splitlines():
        line = line.strip()
        if not line:
            continue

        # New adapter section heuristic
        if line.endswith(":") and ("adapter" in line.lower() or "ethernet" in line.lower() or "wi-fi" in line.lower() or "wireless" in line.lower()):
            if current:
                adapters.append(current)
            current = {"name": line.rstrip(":")}
            continue

        if "IPv4 Address" in line or "IPv4-adresse" in line or "IPv4-adres" in line:
            m = re.search(r"(\d+\.\d+\.\d+\.\d+)", line)
            if m:
                current["ipv4"] = m.group(1)
        if "Subnet Mask" in line or "Sous-réseau" in line or "Subnettmaske" in line:
            m = re.search(r"(\d+\.\d+\.\d+\.\d+)", line)
            if m:
                current["mask"] = m.group(1)

    if current:
        adapters.append(current)

    # Choose the first adapter with IPv4 + mask that looks private
    for a in adapters:
        ip = a.get("ipv4")
        mask = a.get("mask")
        if ip and mask:
            try:
                ip_obj = ipaddress.ip_address(ip)
                if ip_obj.is_private:
                    # build network
                    net = ipaddress.IPv4Network(f"{ip}/{mask}", strict=False)
                    # Prefer /24 if mask is bigger, to keep scans short
                    prefix = net.prefixlen
                    if prefix < 24:
                        # narrow to /24 that contains the IP
                        # derive network base for /24 containing IP
                        parts = ip.split(".")
                        cidr = f"{parts[0]}.{parts[1]}.{parts[2]}.0/24"
                    else:
                        cidr = str(net)
                    return cidr
            except Exception:
                continue
    return None

def auto_network():
    """Return a default network CIDR using OS-specific logic, or None."""
    if platform.system().lower() == "windows":
        return windows_ipconfig_network()
    else:
        # Try 'ip addr' (Linux/mac) as a fallback
        code, out, err = shell("ip addr")
        if code != 0:
            return None
        m = re.search(r"inet\s+(\d+\.\d+\.\d+\.\d+)/(\d+)", out)
        if m:
            ip = m.group(1)
            prefix = m.group(2)
            try:
                net = ipaddress.ip_network(f"{ip}/{prefix}", strict=False)
                # prefer /24
                if net.prefixlen < 24:
                    parts = ip.split(".")
                    return f"{parts[0]}.{parts[1]}.{parts[2]}.0/24"
                return str(net)
            except Exception:
                return None
    return None

def ensure_dirs(outdir):
    scans_dir = os.path.join(outdir, "scans")
    reports_dir = os.path.join(outdir, "reports")
    os.makedirs(scans_dir, exist_ok=True)
    os.makedirs(reports_dir, exist_ok=True)
    return scans_dir, reports_dir

def load_config(config_path):
    with open(config_path, "r", encoding="utf-8") as f:
        return json.load(f)

# -------------------------
# Scanning Routines
# -------------------------

def discover_hosts(scanner, target_cidr, discovery_args):
    """
    Host discovery scan, returns list of alive IPs.
    """
    print(f"[+] Discovery: nmap {discovery_args} {target_cidr}")
    scanner.scan(hosts=target_cidr, arguments=discovery_args)
    alive = []
    for host in scanner.all_hosts():
        state = scanner[host].state()
        if state == "up":
            alive.append(host)
    print(f"[+] Alive hosts: {alive if alive else 'None found'}")
    return alive

def service_scan_host(scanner, host, args):
    """
    Run a service/version scan on a single host with given args.
    Returns the host's scan dict.
    """
    print(f"[+] Service scan: nmap {args} {host}")
    scanner.scan(hosts=host, arguments=args)
    # Return the dict for this host if present
    if host in scanner._scan_result.get('scan', {}):
        return scanner._scan_result['scan'][host]
    # Fallback to top-level scan structure
    return scanner._scan_result.get('scan', {}).get(host, {})

def vuln_scan_host(scanner, host, args):
    """
    Run NSE vuln scripts on a single host.
    """
    print(f"[+] Vuln scan: nmap {args} {host}")
    scanner.scan(hosts=host, arguments=args)
    if host in scanner._scan_result.get('scan', {}):
        return scanner._scan_result['scan'][host]
    return scanner._scan_result.get('scan', {}).get(host, {})

# -------------------------
# Reporting Helpers
# -------------------------

def summarize_scan(host_dict):
    """
    Convert a single host scan dict to a compact summary.
    """
    summary = {
        "ip": host_dict.get("addresses", {}).get("ipv4", ""),
        "hostname": "",
        "state": host_dict.get("status", {}).get("state", ""),
        "ports": []
    }
    # hostname
    hostnames = host_dict.get("hostnames", [])
    if hostnames and isinstance(hostnames, list):
        hn = hostnames[0].get("name", "")
        summary["hostname"] = hn

    # protocols
    for proto in host_dict.get("tcp", {}), host_dict.get("udp", {}):
        if not proto:
            continue
        for port, pdata in sorted(proto.items(), key=lambda x: int(x[0])):
            svc = pdata.get("name", "")
            product = pdata.get("product", "")
            version = pdata.get("version", "")
            extrainfo = pdata.get("extrainfo", "")
            summary["ports"].append({
                "port": int(port),
                "service": svc,
                "product": product,
                "version": version,
                "extrainfo": extrainfo
            })
    return summary

def build_word_report(report_path, title, target_cidr, profile_name, profile_args, alive_hosts, summaries, charts_path=None):
    doc = Document()
    doc.add_heading(title, level=1)

    doc.add_heading("1. Overview", level=2)
    p = doc.add_paragraph()
    p.add_run("Target Network: ").bold = True
    p.add_run(target_cidr)
    p.add_run("\nScan Profile: ").bold = True
    p.add_run(f"{profile_name} ({profile_args})")
    p.add_run("\nDate: ").bold = True
    p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    p.add_run("\nAlive hosts discovered: ").bold = True
    p.add_run(str(len(alive_hosts)))

    doc.add_heading("2. Host Summaries", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "IP Address"
    hdr[1].text = "Hostname"
    hdr[2].text = "Open Ports"
    hdr[3].text = "Services & Versions"

    for s in summaries:
        ip = s["ip"] or "Unknown"
        hn = s["hostname"] or "-"
        ports = ", ".join(str(p["port"]) for p in s["ports"]) if s["ports"] else "-"
        services = "; ".join([
            f"{p['service']} {p['product']} {p['version']}".strip()
            for p in s["ports"]
        ]) if s["ports"] else "-"

        row = table.add_row().cells
        row[0].text = ip
        row[1].text = hn
        row[2].text = ports
        row[3].text = services

    if charts_path and os.path.exists(charts_path):
        doc.add_heading("3. Visuals", level=2)
        doc.add_paragraph("Open ports per host:")
        doc.add_picture(charts_path, width=Inches(6))

    doc.add_heading("4. Notes & Ethical Use", level=2)
    doc.add_paragraph(
        "These scans were performed only on authorized networks for learning purposes. "
        "Always obtain explicit permission before scanning any network you do not own."
    )

    doc.save(report_path)
    return report_path

def make_chart_open_ports_per_host(summaries, out_png):
    if not MATPLOTLIB_OK:
        return None
    labels = []
    counts = []
    for s in summaries:
        label = s["ip"] or s["hostname"] or "host"
        labels.append(label)
        counts.append(len(s["ports"]))
    if not labels:
        return None

    # One plot, no specific colors or styles
    plt.figure()
    plt.bar(labels, counts)
    plt.title("Open Ports Per Host")
    plt.xlabel("Host")
    plt.ylabel("Open Port Count")
    plt.xticks(rotation=45, ha="right")
    plt.tight_layout()
    plt.savefig(out_png, dpi=160)
    plt.close()
    return out_png

# -------------------------
# Main
# -------------------------

def main():
    parser = argparse.ArgumentParser(description="NetScope — Automated Nmap Recon & Reporting")
    parser.add_argument("--config", default=os.path.join(os.path.dirname(__file__), "..", "config", "settings.json"))
    parser.add_argument("--target", help="CIDR to scan (overrides config)")
    parser.add_argument("--profile", choices=["quick", "full", "vuln"], default="quick")
    parser.add_argument("--discovery-only", action="store_true", help="Only run host discovery and exit")
    parser.add_argument("--out", default=os.path.join(os.path.dirname(__file__), ".."), help="Output base folder")
    parser.add_argument("--title", default="NetScope — Nmap Scan Report", help="Title to display in the Word report")
    args = parser.parse_args()

    # Load config
    cfg_path = os.path.abspath(args.config)
    if not os.path.exists(cfg_path):
        print(f"[!] Config not found: {cfg_path}")
        sys.exit(1)

    cfg = load_config(cfg_path)

    # Resolve target CIDR
    target_cidr = args.target
    if not target_cidr:
        if cfg.get("network", "AUTO").upper() == "AUTO":
            target_cidr = auto_network()
            if not target_cidr:
                print("[!] Could not auto-detect network. Set 'network' in config or pass --target.")
                sys.exit(1)
            else:
                print(f"[+] Auto-detected network: {target_cidr}")
        else:
            target_cidr = cfg["network"]
            print(f"[+] Using config network: {target_cidr}")

    # Resolve profile args
    profiles = cfg.get("profiles", {})
    profile_args = profiles.get(args.profile)
    if not profile_args:
        print(f"[!] Unknown profile: {args.profile}")
        sys.exit(1)

    # Prepare output dirs
    outbase = os.path.abspath(args.out)
    scans_dir, reports_dir = ensure_dirs(outbase)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    raw_json_path = os.path.join(scans_dir, f"netscope_{args.profile}_{ts}.json")
    csv_path = os.path.join(scans_dir, f"netscope_{args.profile}_{ts}.csv")
    chart_path = os.path.join(reports_dir, f"open_ports_per_host_{ts}.png")
    report_path = os.path.join(reports_dir, f"NetScope_Report_{ts}.docx")

    # Initialize scanner
    scanner = nmap.PortScanner()

    # Phase 1: Discovery
    alive_hosts = discover_hosts(scanner, target_cidr, cfg.get("discovery_args", "-sn"))
    if not alive_hosts:
        print("[!] No alive hosts found. Exiting.")
        # Still produce an empty report for consistency
        build_word_report(report_path, args.title, target_cidr, args.profile, profile_args, alive_hosts, [])
        print(f"[+] Empty report written: {report_path}")
        sys.exit(0)

    if args.discovery_only:
        print("[+] Discovery-only complete.")
        sys.exit(0)

    # Phase 2/3: Per-host scans
    results = {}
    summaries = []
    for host in alive_hosts:
        if args.profile in ("quick", "full"):
            hdict = service_scan_host(scanner, host, profile_args)
        elif args.profile == "vuln":
            hdict = vuln_scan_host(scanner, host, profile_args)
        else:
            hdict = {}

        if not hdict:
            continue
        results[host] = hdict
        # Enrich IPv4 address (sometimes missing)
        if "addresses" not in hdict:
            hdict["addresses"] = {"ipv4": host}
        summaries.append(summarize_scan(hdict))

    # Save raw json
    with open(raw_json_path, "w", encoding="utf-8") as f:
        json.dump(results, f, indent=2)
    print(f"[+] Raw JSON saved: {raw_json_path}")

    # Save CSV summary
    import csv
    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["ip", "hostname", "port", "service", "product", "version", "extrainfo"])
        for s in summaries:
            if not s["ports"]:
                writer.writerow([s["ip"], s["hostname"], "", "", "", "", ""])
                continue
            for p in s["ports"]:
                writer.writerow([s["ip"], s["hostname"], p["port"], p["service"], p["product"], p["version"], p["extrainfo"]])
    print(f"[+] CSV saved: {csv_path}")

    # Chart (optional)
    if summaries:
        try:
            made = make_chart_open_ports_per_host(summaries, chart_path)
            if made:
                print(f"[+] Chart saved: {made}")
        except Exception as e:
            print(f"[!] Chart error: {e}")

    # Word report
    build_word_report(report_path, args.title, target_cidr, args.profile, profile_args, alive_hosts, summaries, charts_path=chart_path if os.path.exists(chart_path) else None)
    print(f"[+] Report saved: {report_path}")
    print("[+] Done.")

if __name__ == "__main__":
    main()
